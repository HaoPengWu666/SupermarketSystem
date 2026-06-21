from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
PROJECT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "软件工程课程设计报告.docx"
OUTPUT = ROOT / "软件工程课程设计报告-完善版（图片修复）.docx"
SCREENSHOTS = PROJECT / "docs" / "screenshots"


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_run_font(run, name="宋体", size=Pt(10.5), bold=None):
    run.font.name = name
    run.font.size = size
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_p.addnext(new_para._p)
    new_para._p.getparent().remove(new_para._p)
    new_p.getparent().replace(new_p, new_para._p)
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    if text:
        set_run_font(new_para.add_run(text))
    return new_para


def add_text_after(paragraph, lines):
    anchor = paragraph
    for text in lines:
        new_p = insert_paragraph_after(anchor, text, "Normal")
        new_p.paragraph_format.line_spacing = Pt(20)
        new_p.paragraph_format.space_after = Pt(0)
        anchor = new_p
    return anchor


def replace_text(paragraph, text):
    clear_paragraph(paragraph)
    set_run_font(paragraph.add_run(text))


def add_image(paragraph, image_path, alt_text):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 图片段落不能继承正文的“固定值 20 磅”，否则 WPS 会把嵌入型图片裁成一条细线。
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(image_path), width=Cm(14.6))
    doc_pr = shape._inline.docPr
    doc_pr.set("name", alt_text)
    doc_pr.set("descr", alt_text)


doc = Document(SOURCE)

# 删除模板明确标注“定稿后删除”的说明页，保留原封面与主体结构。
for paragraph in list(doc.paragraphs[:15]):
    remove_paragraph(paragraph)

# 按任务书要求，将正文主要样式统一为固定值 20 磅行距。
for style_name in ("Normal", "ds-markdown-paragraph"):
    if style_name in doc.styles:
        style = doc.styles[style_name]
        style.paragraph_format.line_spacing = Pt(20)
        style.paragraph_format.space_after = Pt(0)

# 修正文档中仍保留的协作占位描述，并与真实公开仓库状态保持一致。
for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith("版本控制："):
        replace_text(p, "版本控制：使用 Git 与 GitHub 管理源码，仓库地址：https://github.com/HaoPengWu666/SupermarketSystem")
    elif text.startswith("沟通机制："):
        replace_text(p, "沟通机制：围绕需求、建模、代码和报告进行阶段沟通，重要修改通过提交记录追踪。")
    elif text.startswith("文档协作："):
        replace_text(p, "文档协作：组员负责需求与建模材料整理，组长负责实现内容核对和最终合并。")
    elif text.startswith("任务跟踪："):
        replace_text(p, "任务跟踪：按需求分析、系统设计、编码实现、测试与答辩四阶段分工推进。")
    elif "请在此处插入Git提交记录截图" in text:
        replace_text(p, "GitHub 主分支现有“初始提交”“添加课程设计数据库文件”等提交记录；最终功能与文档修改继续通过独立提交保留变更依据。")

# 插入基于当前 Qt Designer 布局制作的高分辨率界面效果图。
image_map = {
    "插入登录界面截图": ("login.png", "超市销售管理系统登录界面"),
    "插入主界面截图": ("sales.png", "超市销售管理系统主界面"),
    "插入销售界面截图": ("sales.png", "销售收银界面"),
    "插入会员界面截图": ("member.png", "会员管理界面"),
    "插入库存界面截图": ("inventory.png", "库存管理界面"),
    "插入报表界面截图": ("report.png", "经营报表界面"),
}
for p in doc.paragraphs:
    for marker, (filename, alt) in image_map.items():
        if marker in p.text:
            add_image(p, SCREENSHOTS / filename, alt)
            break

# 同时修复模板中原有图片段落，避免任何图片继续继承固定行距。
for p in doc.paragraphs:
    if p._p.xpath(".//w:drawing") or p._p.xpath(".//w:pict"):
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

# 统一现有图题格式。
for p in doc.paragraphs:
    if p.text.strip().startswith("图4."):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = False
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            set_run_font(run, size=Pt(10.5))

# 在实现章节补充本次新增的两类安全删除设计说明。
for p in list(doc.paragraphs):
    if p.text.strip() == "图4.4 会员管理界面":
        add_text_after(p, [
            "会员删除：用户在会员表格中选中记录后点击“删除会员”。界面先进行二次确认，再调用“删除会员”存储过程。若会员已有消费记录，数据库拒绝物理删除并返回明确提示；无消费历史的测试会员可正常删除。",
            "该设计将业务历史完整性检查放在数据库层统一执行，避免不同客户端绕过校验。",
        ])
    elif p.text.strip() == "图4.5 库存管理界面":
        add_text_after(p, [
            "商品删除：库存员选中商品后点击“删除商品”，确认后调用同名存储过程。存储过程同时检查销售明细表和进货明细表；已参与业务的商品保留历史，新建且无业务历史的商品允许删除。",
            "Qt 端删除成功后调用 refresh() 同步刷新库存列表和预警列表；失败时直接显示 MySQL 返回的完整性提示。",
        ])

# 更新数据库组件数量（新增两个安全删除存储过程）。
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if "8张表、7视图、3存储过程、2触发器" in cell.text:
                cell.text = cell.text.replace("3存储过程", "5存储过程")

# 需求表中补充会员安全删除条目。
member_req = doc.tables[8]
if not any("会员删除" in row.cells[1].text for row in member_req.rows[1:]):
    cells = member_req.add_row().cells
    for cell, value in zip(cells, ["FR-M06", "会员删除", "无消费历史的会员可删除；有消费历史时阻止并提示"]):
        cell.text = value

# 数据访问接口表补充删除过程接口。
interface_table = doc.tables[27]
existing_interface_text = " ".join(cell.text for row in interface_table.rows for cell in row.cells)
if "deleteProduct" not in existing_interface_text:
    for values in [
        ["deleteProduct", "barcode", "bool/错误信息", "调用删除商品存储过程"],
        ["deleteMember", "memberId", "bool/错误信息", "调用删除会员存储过程"],
    ]:
        cells = interface_table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value

# 测试用例增加正常删除与历史数据保护场景；不虚构本环境未执行的动态联调结果。
test_table = doc.tables[29]
existing_test_text = " ".join(cell.text for row in test_table.rows for cell in row.cells)
if "TC-18" not in existing_test_text:
    new_tests = [
        ["TC-18", "库存", "删除无历史商品", "新增测试商品后删除", "确认后删除成功并刷新列表", "代码与SQL静态检查一致", "待目标机联调"],
        ["TC-19", "库存", "保护历史商品", "删除已有进货记录商品", "拒绝删除并提示历史数据保护", "代码与SQL静态检查一致", "待目标机联调"],
        ["TC-20", "会员", "删除无消费会员", "新增测试会员后删除", "确认后删除成功并刷新列表", "代码与SQL静态检查一致", "待目标机联调"],
        ["TC-21", "会员", "保护消费会员", "删除已有消费记录会员", "拒绝删除并提示历史数据保护", "代码与SQL静态检查一致", "待目标机联调"],
    ]
    for values in new_tests:
        cells = test_table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value

# 修正测试汇总，区分历史用例与新增的待目标机动态验证用例。
for table in doc.tables:
    for row in table.rows:
        if len(row.cells) >= 4 and row.cells[0].text.strip() == "功能测试":
            row.cells[1].text = "21"
            row.cells[2].text = "17（另4项待目标机联调）"
            row.cells[3].text = "既有用例100%"

# 为所有新增表格行统一基本字体、行距和垂直居中。
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=Pt(9))

# 清理残余截图占位文本。
for p in doc.paragraphs:
    if "【请在此处插入" in p.text:
        replace_text(p, "")

doc.save(OUTPUT)
print(OUTPUT)
